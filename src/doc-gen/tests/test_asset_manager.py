import os
from pathlib import Path
import asyncio
import pytest

import asset_manager as am


def test_move_assets_and_summary(tmp_path):
    # Create source files to act as assets
    src1 = tmp_path / "file1.txt"
    src1.write_text("hello world")

    nested = tmp_path / "nested"
    nested.mkdir()
    src2 = nested / "doc.md"
    src2.write_text("# Title\n\nSome content")

    job_folder = tmp_path / "job"

    # If required libs missing, AssetManager should raise ImportError
    if not am.HAS_PDF or not am.HAS_DOCX:
        with pytest.raises(ImportError):
            _ = am.AssetManager(job_id="testjob", job_folder=str(job_folder), assets=[str(src1), str(src2)])
        return

    manager = am.AssetManager(job_id="testjob", job_folder=str(job_folder), assets=[str(src1), str(src2)])

    # Ensure files were moved/copied
    assert len(manager.moved_files) == 2
    for p in manager.moved_files:
        assert Path(p).exists()

    summary = manager.get_asset_summary()
    assert summary["total_files"] == 2
    assert isinstance(summary["types"], dict)

    files = manager.list_asset_files()
    assert len(files) == 2

    formatted = manager.format_assets_for_agent()
    assert "Total files" in formatted


def _make_fake_memory_class():
    class FakeMemory:
        def __init__(self, config=None):
            self.added = []

        async def add(self, mem):
            # store content string or MemoryContent-like object
            self.added.append(mem)

    return FakeMemory


def test_create_vector_memory_skips_when_unavailable(monkeypatch, tmp_path):
    # Force the module to behave as if AutoGen memory is unavailable
    monkeypatch.setattr(am, "HAS_AUTOGEN_MEMORY", False)
    # AssetManager now requires both PDF and DOCX libs; if they are missing,
    # constructing the manager should raise ImportError.
    if not (am.HAS_PDF and am.HAS_DOCX):
        with pytest.raises(ImportError):
            _ = am.AssetManager(job_id="j2", job_folder=str(tmp_path / "job2"), assets=[])
        return

    manager = am.AssetManager(job_id="j2", job_folder=str(tmp_path / "job2"), assets=[])

    memory = asyncio.run(manager.create_vector_memory())
    assert memory is None


def test_docx_processing_and_vector_memory(monkeypatch, tmp_path):
    # Require both DOCX and PDF libs for AssetManager initialization
    if not (am.HAS_DOCX and am.HAS_PDF):
        pytest.skip("python-docx or pypdf not installed; skipping DOCX processing integration test")

    # Create a simple DOCX file using the real library
    from docx import Document as DocxDocument

    doc_path = tmp_path / "sample.docx"
    d = DocxDocument()
    d.add_paragraph("Hello from paragraph one.")
    d.add_paragraph("Another paragraph with more text.")
    d.save(str(doc_path))

    # Provide a fake memory implementation and enable autogen path
    FakeMemory = _make_fake_memory_class()
    monkeypatch.setattr(am, "HAS_AUTOGEN_MEMORY", True)
    monkeypatch.setattr(am, "ChromaDBVectorMemory", FakeMemory)
    # Provide lightweight config classes the constructor expects
    monkeypatch.setattr(am, "PersistentChromaDBVectorMemoryConfig", lambda **kwargs: kwargs)
    monkeypatch.setattr(am, "SentenceTransformerEmbeddingFunctionConfig", lambda **kwargs: kwargs)

    manager = am.AssetManager(job_id="docxjob", job_folder=str(tmp_path / "job_docx"), assets=[str(doc_path)])
    memory = asyncio.run(manager.create_vector_memory())

    assert memory is not None
    # FakeMemory stores added items in .added
    assert hasattr(memory, "added")
    assert len(memory.added) > 0


def test_pdf_processing_and_vector_memory_with_fake_reader(monkeypatch, tmp_path):
    # Require both DOCX and PDF libs for AssetManager initialization
    if not (am.HAS_PDF and am.HAS_DOCX):
        pytest.skip("pypdf or python-docx not installed; skipping PDF integration test")

    # Create a fake PdfReader that returns pages with extract_text
    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        def __init__(self, f):
            # ignore f, provide two fake pages
            self.pages = [FakePage("Page one text"), FakePage("Page two text")]

    # Monkeypatch the PdfReader in the module
    monkeypatch.setattr(am, "PdfReader", FakeReader)

    # Provide fake memory as above
    FakeMemory = _make_fake_memory_class()
    monkeypatch.setattr(am, "HAS_AUTOGEN_MEMORY", True)
    monkeypatch.setattr(am, "ChromaDBVectorMemory", FakeMemory)
    monkeypatch.setattr(am, "PersistentChromaDBVectorMemoryConfig", lambda **kwargs: kwargs)
    monkeypatch.setattr(am, "SentenceTransformerEmbeddingFunctionConfig", lambda **kwargs: kwargs)

    # Create a dummy pdf file (empty file content is fine since FakeReader ignores it)
    pdf_path = tmp_path / "doc.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%EOF\n")

    manager = am.AssetManager(job_id="pdfjob", job_folder=str(tmp_path / "job_pdf"), assets=[str(pdf_path)])
    memory = asyncio.run(manager.create_vector_memory())

    assert memory is not None
    assert hasattr(memory, "added")
    assert len(memory.added) >= 2
